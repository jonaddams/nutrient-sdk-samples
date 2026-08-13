"""
Generate the sample DOCX for the .NET SDK "Document Fonts" sample.

The document deliberately requests two fonts that are NOT installed on the
machine rendering it:

  - Alfa Slab One (headings) -- a heavy slab serif. Substitution is obvious:
    weight, width and line breaks all change at once.
  - EB Garamond (body) -- a text serif. Substitution here is subtler, shifting
    metrics and re-flowing paragraphs.

Both are SIL OFL licensed and ship in public/fonts/ for the viewer to load as
customFonts. Do NOT install them system-wide: the whole point is that the
renderer lacks them until they are supplied explicitly.

Usage: python3 scripts/generate-acme-sow.py
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

HEADING_FONT = "Alfa Slab One"
BODY_FONT = "EB Garamond"
OUT = "public/documents/dotnet-sdk/acme-sow.docx"


def styled(paragraph, text, font, size, bold=False):
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    return run


def main():
    doc = Document()
    for section in doc.sections:
        section.left_margin = section.right_margin = Inches(1)

    styled(doc.add_paragraph(), "ACME STATEMENT OF WORK", HEADING_FONT, 26)

    meta = doc.add_paragraph()
    styled(meta, "Prepared for: Northwind Ltd\nDate: 2026-08-12\n"
                 "Reference: SOW-2026-0142", BODY_FONT, 11)

    styled(doc.add_paragraph(), "1. Scope of Services", HEADING_FONT, 16)
    styled(
        doc.add_paragraph(),
        "Acme will design, implement and support the document processing "
        "pipeline described in Schedule A. Work is delivered in three phases, "
        "each concluding with a written acceptance review. Any change to the "
        "agreed scope is handled through the change-control process in "
        "Section 4 and requires written approval from both parties.",
        BODY_FONT, 11)

    styled(doc.add_paragraph(), "2. Fees", HEADING_FONT, 16)

    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    data = [
        ("Phase", "Description", "Fee"),
        ("1", "Discovery and solution design", "$24,000"),
        ("2", "Implementation and integration", "$86,500"),
        ("3", "Rollout, training and handover", "$18,750"),
    ]
    for row, values in zip(table.rows, data):
        for cell, value in zip(row.cells, values):
            cell.text = ""
            styled(cell.paragraphs[0], value, BODY_FONT, 10,
                   bold=(values is data[0]))

    total = doc.add_paragraph()
    total.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    styled(total, "Total: $129,250", HEADING_FONT, 13)

    styled(doc.add_paragraph(), "3. Acceptance", HEADING_FONT, 16)
    styled(
        doc.add_paragraph(),
        "Signed for and on behalf of the parties below.",
        BODY_FONT, 11)

    sign = doc.add_paragraph()
    styled(sign, "\n\n_______________________________        "
                 "_______________________________\n"
                 "Acme Corp                                            "
                 "Northwind Ltd", BODY_FONT, 10)

    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
